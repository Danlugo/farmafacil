"""API route definitions."""

import csv
import io
import logging
import re
from html import escape
from uuid import uuid4

import secrets

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from sqlalchemy import func, select, text

from farmafacil import __version__
from farmafacil.api.limiter import limiter
from farmafacil.db.session import async_session
from typing import Literal

from pydantic import BaseModel, Field

from farmafacil.config import ADMIN_PASSWORD, ADMIN_USERNAME
from farmafacil.bot.handler import handle_incoming_message
from farmafacil.bot.whatsapp import start_collecting, stop_collecting
from farmafacil.models.database import ConversationLog, IntentKeyword, SearchLog, User, VoiceMessage
from farmafacil.services.conversation_log import log_inbound, log_outbound
from farmafacil.models.schemas import HealthResponse, SearchRequest, SearchResponse
from farmafacil.services.search import search_drug

router = APIRouter()
_http_basic = HTTPBasic()
logger = logging.getLogger(__name__)


def _require_admin(
    credentials: HTTPBasicCredentials = Depends(_http_basic),
) -> str:
    """Verify HTTP Basic credentials match the admin user.

    Used to protect endpoints that serve sensitive data (e.g. user audio)
    outside the SQLAdmin session-cookie scope.
    """
    correct_user = secrets.compare_digest(
        credentials.username.encode(), ADMIN_USERNAME.encode(),
    )
    correct_pass = secrets.compare_digest(
        credentials.password.encode(), ADMIN_PASSWORD.encode(),
    )
    if not (correct_user and correct_pass):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return credentials.username


@router.get("/health", response_model=None)
async def health_check() -> HealthResponse | JSONResponse:
    """Health check endpoint with DB connectivity probe.

    Executes ``SELECT 1`` against the application database.  Returns HTTP 200
    with ``{"status": "ok", ...}`` on success, or HTTP 503 with
    ``{"status": "unhealthy", ...}`` when the database is unreachable.

    Docker and uptime monitors rely on this endpoint — the ``status`` key
    is preserved for backwards compatibility.

    (Item 82, v0.25.0 — added DB probe.)
    """
    try:
        async with async_session() as session:
            await session.execute(text("SELECT 1"))
        return HealthResponse(status="ok", version=__version__)
    except Exception:
        logger.error("Health check DB probe failed", exc_info=True)
        return JSONResponse(
            content={"status": "unhealthy", "version": __version__},
            status_code=503,
        )


@router.post("/api/v1/search", response_model=SearchResponse)
@limiter.limit("30/minute")
async def search(request: Request, body: SearchRequest) -> SearchResponse:
    """Search for a drug across all pharmacies."""
    return await search_drug(body.query, city=body.city)


@router.get("/api/v1/search", response_model=SearchResponse)
@limiter.limit("30/minute")
async def search_get(
    request: Request,
    q: str = Query(..., min_length=2, max_length=200),
    city: str | None = Query(None, max_length=50),
) -> SearchResponse:
    """Search for a drug via GET (convenience for WhatsApp bot / browser)."""
    return await search_drug(q, city=city)


class ChatRequest(BaseModel):
    """Chat API request — mirrors the WhatsApp message flow for external bots."""

    sender_id: str = Field(
        ..., min_length=5, max_length=30,
        description="Phone number identifying the user (e.g. '584127006823')",
    )
    sender_name: str = Field(
        "", max_length=100,
        description="Display name of the sender — used to pre-fill the user's "
        "name during onboarding so the bot can greet them by name.",
    )
    text: str = Field(
        ..., min_length=1, max_length=2000,
        description="Message text to process",
    )


class ChatResponseItem(BaseModel):
    """A single response item (text, image, or interactive list)."""

    type: Literal["text", "image", "list"] = Field(
        ..., description="Response type: text, image, or list",
    )
    body: str | None = Field(None, description="Text body (for text and list types)")
    url: str | None = Field(None, description="Image URL (for image type)")
    caption: str | None = Field(None, description="Image caption (for image type)")
    button: str | None = Field(None, description="Button label (for list type)")
    rows: list[dict] | None = Field(None, description="List rows (for list type)")
    header: str | None = Field(None, description="Header text (for list type)")
    footer: str | None = Field(None, description="Footer text (for list type)")


class ChatResponse(BaseModel):
    """Chat API response — contains all messages the bot would have sent."""

    responses: list[ChatResponseItem]


async def _log_relay_responses(phone: str, collected: list[dict]) -> None:
    """Log outbound relay responses for conversation context.

    Best-effort: failures are logged but never propagated to the caller.
    Each collected item's ``body`` or ``caption`` is stored as an outbound
    ``text`` entry in ``conversation_log`` so that ``get_recent_history()``
    can provide AI context for follow-up questions from relay users.

    Voice-ack messages (``🎙️ Te escuché: …``) are skipped because they
    just echo the transcription that is already logged as inbound — keeping
    them would add noise to the AI classifier's context window.

    .. note:: Each outbound item is a separate ``INSERT + COMMIT``.  At
       current traffic this is fine; if response counts grow significantly,
       consider a bulk-insert helper in ``conversation_log``.
    """
    for item in collected:
        text = item.get("body") or item.get("caption") or ""
        if not text.strip():
            continue
        # Skip voice-ack echo — already captured as inbound transcription
        if text.startswith("🎙️ Te escuché"):
            continue
        try:
            await log_outbound(phone, text)
        except Exception:
            logger.warning(
                "relay: failed to log outbound for %s", phone, exc_info=True,
            )


@router.post("/api/v1/chat", response_model=ChatResponse)
@limiter.limit("120/minute")
async def chat(request: Request, body: ChatRequest) -> ChatResponse:
    """Process a text message through the full FarmaFacil handler.

    Runs the exact same logic as a direct WhatsApp message — intent
    detection, drug search, onboarding, feedback, help — but returns
    the bot's responses as JSON instead of sending them via WhatsApp.

    Designed for relay bots (e.g. Chamo in a WhatsApp group) that
    forward group messages to FarmaFacil and post the responses back.

    Intentionally unauthenticated — callers are rate-limited and the
    endpoint runs the same handler as the webhook, which has its own
    input validation.  Chamo connects via localhost on the same server.

    Args:
        body: ChatRequest with sender_id, sender_name, and text.

    Returns:
        ChatResponse with an ordered list of response items.
    """
    # Log inbound for conversation context — the AI classifier in
    # handle_incoming_message calls get_recent_history() which needs
    # prior messages to understand follow-up questions like "which is
    # the cheapest?" after a drug search.
    try:
        await log_inbound(body.sender_id, body.text)
    except Exception:
        logger.warning(
            "chat: failed to log inbound for %s", body.sender_id, exc_info=True,
        )

    # Enter proxy mode: outbound send_* calls collect into a list
    try:
        start_collecting()
        await handle_incoming_message(
            sender=body.sender_id,
            message_text=body.text,
            wa_profile_name=body.sender_name,
        )
    except Exception:
        logger.error(
            "Chat handler failed for sender=%s text=%r",
            body.sender_id, body.text[:100], exc_info=True,
        )
    finally:
        collected = stop_collecting()

    # Log outbound for conversation context (best-effort)
    await _log_relay_responses(body.sender_id, collected)

    return ChatResponse(responses=[ChatResponseItem(**item) for item in collected])


@router.post("/api/v1/chat/voice", response_model=ChatResponse)
@limiter.limit("30/minute")
async def chat_voice(
    request: Request,
    sender_id: str = Form(..., min_length=5, max_length=30),
    sender_name: str = Form("", max_length=100),
    audio: UploadFile = File(...),
) -> ChatResponse:
    """Process a voice message through the full FarmaFacil handler.

    Accepts a raw audio file upload from relay bots (e.g. Chamo) that can
    download group voice notes but cannot supply a WhatsApp Media API
    ``media_id``.  The audio is saved to disk, transcribed via Whisper,
    and fed into ``handle_incoming_message`` exactly as a native WhatsApp
    voice note would be — including a ``VoiceMessage`` DB record and the
    ``voice_message_id`` linkage through to any resulting search or
    feedback log.

    The endpoint runs in proxy mode: all outbound ``send_*`` calls are
    intercepted and returned as a ``ChatResponse`` instead of being sent
    via WhatsApp.

    Intentionally unauthenticated — callers are rate-limited (30/min,
    lower than the text endpoint due to Whisper API cost) and Chamo
    connects via localhost on the same server.  In production, Docker
    Compose binds the app port to ``127.0.0.1`` only (docker-compose.yml),
    so only same-host processes can reach this endpoint.

    Args:
        sender_id: Phone number identifying the user (e.g. ``'584127006823'``).
        sender_name: Display name of the sender — used to pre-fill the
            user's name during onboarding.
        audio: The audio file bytes (OGG, MP3, M4A, etc.).

    Returns:
        ChatResponse with an ordered list of response items.

    Raises:
        HTTPException 413: Audio exceeds the 25 MB Whisper limit.
    """
    from farmafacil.services.voice import (
        MAX_AUDIO_BYTES,
        get_audio_absolute_path,
        save_audio_file,
        transcribe_audio,
    )
    from farmafacil.services.users import get_or_create_user, validate_user_profile

    # --- 1. Read and size-check the upload --------------------------------
    # NOTE: The entire file is buffered before the size check.  With the
    # 30/min rate limit and max concurrency of ~1–2 requests, peak memory
    # for this path is ~50 MB which is acceptable.  If traffic grows,
    # switch to chunked reading with an early abort.
    audio_data = await audio.read()
    if len(audio_data) > MAX_AUDIO_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"Audio file exceeds the {MAX_AUDIO_BYTES // (1024 * 1024)} MB limit.",
        )

    # --- 2. Get or create the user ----------------------------------------
    try:
        user = await get_or_create_user(sender_id, wa_profile_name=sender_name)
        user = await validate_user_profile(user)
    except Exception:
        logger.error(
            "chat_voice: user lookup failed for sender=%s", sender_id, exc_info=True,
        )
        raise HTTPException(status_code=500, detail="Error al procesar el usuario.")

    # --- 3. Save audio to disk --------------------------------------------
    relay_id = f"relay_{uuid4().hex[:12]}"
    audio_relative_path = save_audio_file(audio_data, user.id, relay_id)
    audio_absolute_path = get_audio_absolute_path(audio_relative_path)

    # --- 4. Transcribe via Whisper ----------------------------------------
    transcription, detected_lang, duration = await transcribe_audio(audio_absolute_path)

    # --- 5. Persist VoiceMessage record -----------------------------------
    voice_msg_id: int | None = None
    try:
        async with async_session() as session:
            voice_msg = VoiceMessage(
                user_id=user.id,
                phone_number=sender_id,
                audio_path=audio_relative_path,
                duration_seconds=duration,
                original_language=detected_lang,
                transcription=transcription,
                wa_message_id=relay_id,
                conversation_log_id=None,
                transcription_model="whisper-1" if transcription else None,
            )
            session.add(voice_msg)
            await session.commit()
            await session.refresh(voice_msg)
            voice_msg_id = voice_msg.id
            logger.info(
                "chat_voice: VoiceMessage saved id=%d user=%d duration=%.1fs transcription='%s'",
                voice_msg_id, user.id, duration or 0, (transcription or "")[:60],
            )
    except Exception:
        logger.error(
            "chat_voice: failed to save VoiceMessage for sender=%s relay_id=%s",
            sender_id, relay_id, exc_info=True,
        )
        # Orphan audio cleanup — best effort
        try:
            audio_absolute_path.unlink(missing_ok=True)
        except OSError:
            logger.warning("chat_voice: could not remove orphan audio: %s", audio_absolute_path)
        raise HTTPException(status_code=500, detail="Error al guardar el mensaje de voz.")

    # --- 6. Transcription failed — return failure response ----------------
    if not transcription:
        logger.info(
            "chat_voice: transcription empty for sender=%s relay_id=%s", sender_id, relay_id,
        )
        # Clean up the audio file — the VoiceMessage record stays for audit
        # but there is no value keeping the raw bytes if Whisper couldn't decode.
        try:
            audio_absolute_path.unlink(missing_ok=True)
        except OSError:
            logger.warning("chat_voice: could not remove audio for empty transcription: %s", audio_absolute_path)
        return ChatResponse(
            responses=[ChatResponseItem(type="text", body="🎙️ No pude entender el audio.")]
        )

    # --- 7. Log inbound transcription for conversation context ------------
    try:
        await log_inbound(sender_id, transcription)
    except Exception:
        logger.warning(
            "chat_voice: failed to log inbound for %s", sender_id, exc_info=True,
        )

    # --- 8. Run handler in proxy mode -------------------------------------
    collected: list[dict] = []
    try:
        start_collecting()
        # Acknowledgment mirrors handle_voice_message in handler.py
        from farmafacil.bot.whatsapp import send_text_message
        display_text = transcription[:100] + ("..." if len(transcription) > 100 else "")
        await send_text_message(sender_id, f"🎙️ Te escuché: _{display_text}_")
        await handle_incoming_message(
            sender=sender_id,
            message_text=transcription,
            voice_message_id=voice_msg_id,
            wa_profile_name=sender_name,
        )
    except Exception:
        logger.error(
            "chat_voice: handler failed for sender=%s transcription=%r",
            sender_id, transcription[:100], exc_info=True,
        )
    finally:
        collected = stop_collecting()

    # --- 9. Log outbound for conversation context (best-effort) -----------
    await _log_relay_responses(sender_id, collected)

    return ChatResponse(responses=[ChatResponseItem(**item) for item in collected])


@router.get("/api/v1/conversations")
@limiter.limit("60/minute")
async def get_conversations(
    request: Request,
    phone: str | None = Query(None, max_length=30),
    limit: int = Query(50, le=200),
    _admin: str = Depends(_require_admin),
) -> list[dict]:
    """View conversation logs for troubleshooting.

    Args:
        phone: Optional phone number filter.
        limit: Max records to return (default 50).

    Returns:
        List of conversation log entries.
    """
    async with async_session() as session:
        query = select(ConversationLog).order_by(ConversationLog.created_at.desc())
        if phone:
            query = query.where(ConversationLog.phone_number.contains(phone))
        query = query.limit(limit)
        result = await session.execute(query)
        logs = result.scalars().all()

        return [
            {
                "id": log.id,
                "phone": log.phone_number,
                "direction": log.direction,
                "message": log.message_text[:500],
                "type": log.message_type,
                "wa_id": log.wa_message_id,
                "time": log.created_at.isoformat() if log.created_at else None,
            }
            for log in logs
        ]


@router.get("/api/v1/users")
@limiter.limit("60/minute")
async def get_users(
    request: Request,
    limit: int = Query(50, le=200),
    _admin: str = Depends(_require_admin),
) -> list[dict]:
    """View registered users.

    Args:
        limit: Max records to return.

    Returns:
        List of user records.
    """
    async with async_session() as session:
        result = await session.execute(
            select(User).order_by(User.created_at.desc()).limit(limit)
        )
        users = result.scalars().all()

        return [
            {
                "id": u.id,
                "phone": u.phone_number,
                "name": u.name,
                "zone": u.zone_name,
                "city_code": u.city_code,
                "lat": u.latitude,
                "lng": u.longitude,
                "display_preference": u.display_preference,
                "onboarding_step": u.onboarding_step,
                "created": u.created_at.isoformat() if u.created_at else None,
            }
            for u in users
        ]


# ── Intent Keywords Management ──────────────────────────────────────────


class IntentCreate(BaseModel):
    action: str = Field(..., min_length=1, max_length=50)
    keyword: str = Field(..., min_length=1, max_length=100)
    response: str | None = Field(None, max_length=2000)


@router.get("/api/v1/intents")
@limiter.limit("30/minute")
async def get_intents(
    request: Request,
    action: str | None = Query(None, max_length=50),
    _admin: str = Depends(_require_admin),
) -> list[dict]:
    """List all intent keywords, optionally filtered by action."""
    async with async_session() as session:
        query = select(IntentKeyword).order_by(IntentKeyword.action, IntentKeyword.keyword)
        if action:
            query = query.where(IntentKeyword.action == action)
        result = await session.execute(query)
        intents = result.scalars().all()
        return [
            {
                "id": i.id,
                "action": i.action,
                "keyword": i.keyword,
                "response": i.response,
                "is_active": i.is_active,
            }
            for i in intents
        ]


@router.post("/api/v1/intents")
@limiter.limit("30/minute")
async def create_intent(request: Request, data: IntentCreate, _admin: str = Depends(_require_admin)) -> dict:
    """Add a new intent keyword."""
    async with async_session() as session:
        intent = IntentKeyword(
            action=data.action,
            keyword=data.keyword.lower().strip(),
            response=data.response,
            is_active=True,
        )
        session.add(intent)
        await session.commit()
        await session.refresh(intent)
        # Invalidate cache
        from farmafacil.services.intent import _load_keyword_cache
        await _load_keyword_cache()
        return {"id": intent.id, "action": intent.action, "keyword": intent.keyword}


@router.get("/api/v1/stats")
@limiter.limit("60/minute")
async def get_stats(
    request: Request,
    phone: str | None = Query(None, max_length=30),
    _admin: str = Depends(_require_admin),
) -> dict:
    """Usage statistics — global totals or per-user breakdown.

    Args:
        phone: Optional phone number to get per-user stats.

    Returns:
        Dict with questions, tokens, and success counts.
    """
    async with async_session() as session:
        if phone:
            result = await session.execute(
                select(User).where(User.phone_number == phone)
            )
            user = result.scalar_one_or_none()
            if not user:
                return {"error": "user not found"}
            from farmafacil.services.chat_debug import get_user_stats

            stats = await get_user_stats(phone, user.id)
            return {"phone": phone, "name": user.name, **stats}

        # Global stats
        total_users = (
            await session.execute(select(func.count(User.id)))
        ).scalar() or 0
        total_questions = (
            await session.execute(
                select(func.count(ConversationLog.id)).where(
                    ConversationLog.direction == "inbound"
                )
            )
        ).scalar() or 0
        total_success = (
            await session.execute(
                select(func.count(SearchLog.id)).where(
                    SearchLog.feedback == "yes"
                )
            )
        ).scalar() or 0
        tokens = (
            await session.execute(
                select(
                    func.coalesce(func.sum(User.total_tokens_in), 0),
                    func.coalesce(func.sum(User.total_tokens_out), 0),
                    func.coalesce(func.sum(User.tokens_in_haiku), 0),
                    func.coalesce(func.sum(User.tokens_out_haiku), 0),
                    func.coalesce(func.sum(User.calls_haiku), 0),
                    func.coalesce(func.sum(User.tokens_in_sonnet), 0),
                    func.coalesce(func.sum(User.tokens_out_sonnet), 0),
                    func.coalesce(func.sum(User.calls_sonnet), 0),
                    func.coalesce(func.sum(User.tokens_in_admin), 0),
                    func.coalesce(func.sum(User.tokens_out_admin), 0),
                    func.coalesce(func.sum(User.calls_admin), 0),
                )
            )
        ).one()

        from farmafacil.services.chat_debug import estimate_cost

        cost_haiku = estimate_cost(tokens[2], tokens[3], "haiku")
        cost_sonnet = estimate_cost(tokens[5], tokens[6], "sonnet")
        cost_admin = estimate_cost(tokens[8], tokens[9], "opus")

        return {
            "total_users": total_users,
            "total_questions": total_questions,
            "total_success": total_success,
            "total_tokens_in": tokens[0],
            "total_tokens_out": tokens[1],
            "haiku": {
                "tokens_in": tokens[2],
                "tokens_out": tokens[3],
                "calls": tokens[4],
                "est_cost_usd": round(cost_haiku, 4),
            },
            "sonnet": {
                "tokens_in": tokens[5],
                "tokens_out": tokens[6],
                "calls": tokens[7],
                "est_cost_usd": round(cost_sonnet, 4),
            },
            "admin": {
                "tokens_in": tokens[8],
                "tokens_out": tokens[9],
                "calls": tokens[10],
                "est_cost_usd": round(cost_admin, 4),
            },
            "est_cost_total_usd": round(
                cost_haiku + cost_sonnet + cost_admin, 4
            ),
        }


@router.delete("/api/v1/intents/{intent_id}")
@limiter.limit("30/minute")
async def delete_intent(request: Request, intent_id: int, _admin: str = Depends(_require_admin)) -> dict:
    """Deactivate an intent keyword."""
    async with async_session() as session:
        result = await session.execute(
            select(IntentKeyword).where(IntentKeyword.id == intent_id)
        )
        intent = result.scalar_one_or_none()
        if not intent:
            return {"error": "not found"}
        intent.is_active = False
        await session.commit()
        from farmafacil.services.intent import _load_keyword_cache
        await _load_keyword_cache()
        return {"id": intent.id, "deactivated": True}


# ── Admin User Stats Page ──────────────────────────────────────────────


@router.get("/admin/user-stats/{user_id}", response_class=HTMLResponse)
@limiter.limit("60/minute")
async def admin_user_stats(request: Request, user_id: int, _admin: str = Depends(_require_admin)) -> HTMLResponse:
    """Render an HTML stats dashboard for a single user.

    Args:
        user_id: Database ID of the user.

    Returns:
        HTML page with usage stats, cost estimates, and activity metrics.
    """
    from farmafacil.services.chat_debug import estimate_cost_breakdown, get_user_stats

    async with async_session() as session:
        result = await session.execute(select(User).where(User.id == user_id))
        user = result.scalar_one_or_none()
        if not user:
            return HTMLResponse("<h1>User not found</h1>", status_code=404)

        stats = await get_user_stats(user.phone_number, user.id)

        # Search counts
        total_searches = (
            await session.execute(
                select(func.count(SearchLog.id)).where(SearchLog.user_id == user_id)
            )
        ).scalar() or 0
        successful_searches = stats["total_success"]

        # Recent searches
        recent = await session.execute(
            select(SearchLog.query, SearchLog.results_count, SearchLog.feedback, SearchLog.searched_at)
            .where(SearchLog.user_id == user_id)
            .order_by(SearchLog.searched_at.desc())
            .limit(10)
        )
        recent_searches = recent.all()

    costs = estimate_cost_breakdown(stats)
    success_rate = (successful_searches / total_searches * 100) if total_searches > 0 else 0

    # All user-sourced values are HTML-escaped to prevent stored XSS
    # (names and search queries come from WhatsApp and must not be trusted).
    safe_name = escape(user.name or "Unknown")
    safe_phone = escape(user.phone_number or "")
    safe_zone = escape(user.zone_name or "—")
    safe_city = escape(user.city_code or "—")
    safe_title_name = escape(user.name or user.phone_number or "")

    # Build HTML
    searches_html = ""
    for s in recent_searches:
        fb = s.feedback or "—"
        ts = s.searched_at.strftime("%Y-%m-%d %H:%M") if s.searched_at else "—"
        fb_class = "success" if fb == "yes" else ("danger" if fb == "no" else "")
        searches_html += (
            f"<tr><td>{escape(s.query or '')}</td><td>{s.results_count}</td>"
            f'<td class="{fb_class}">{escape(fb)}</td><td>{ts}</td></tr>'
        )

    html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Stats — {safe_title_name}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               max-width: 900px; margin: 40px auto; padding: 0 20px; color: #333; }}
        h1 {{ color: #1a73e8; }}
        .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 16px; margin: 24px 0; }}
        .card {{ background: #f8f9fa; border-radius: 8px; padding: 20px; border: 1px solid #e0e0e0; }}
        .card .label {{ font-size: 13px; color: #666; text-transform: uppercase; letter-spacing: 0.5px; }}
        .card .value {{ font-size: 28px; font-weight: 700; margin-top: 4px; }}
        .card .sub {{ font-size: 12px; color: #888; margin-top: 4px; }}
        .cost {{ color: #1a73e8; }}
        .success {{ color: #0d904f; }}
        .danger {{ color: #d93025; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 16px; }}
        th, td {{ text-align: left; padding: 10px 12px; border-bottom: 1px solid #e0e0e0; }}
        th {{ background: #f1f3f4; font-size: 13px; text-transform: uppercase; letter-spacing: 0.5px; }}
        a.back {{ display: inline-block; margin-bottom: 16px; color: #1a73e8; text-decoration: none; }}
        a.back:hover {{ text-decoration: underline; }}
        .section {{ margin-top: 32px; }}
        .section h2 {{ font-size: 18px; color: #444; border-bottom: 2px solid #1a73e8; padding-bottom: 8px; }}
    </style>
</head>
<body>
    <a class="back" href="/admin/user/details/{user_id}">&larr; Back to User</a>
    <h1>{safe_name} &mdash; Usage Stats</h1>
    <p>Phone: {safe_phone} &bull; Zone: {safe_zone} &bull; City: {safe_city}</p>

    <div class="section">
        <h2>Activity</h2>
        <div class="grid">
            <div class="card">
                <div class="label">Questions</div>
                <div class="value">{stats['total_questions']}</div>
                <div class="sub">Total inbound messages</div>
            </div>
            <div class="card">
                <div class="label">Searches</div>
                <div class="value">{total_searches}</div>
                <div class="sub">{successful_searches} successful ({success_rate:.0f}%)</div>
            </div>
            <div class="card">
                <div class="label">Success Rate</div>
                <div class="value {'success' if success_rate >= 50 else 'danger'}">{success_rate:.0f}%</div>
                <div class="sub">Positive feedback / total searches</div>
            </div>
        </div>
    </div>

    <div class="section">
        <h2>Token Usage</h2>
        <div class="grid">
            <div class="card">
                <div class="label">Total Tokens</div>
                <div class="value">{stats['total_tokens_in'] + stats['total_tokens_out']:,}</div>
                <div class="sub">{stats['total_tokens_in']:,} in / {stats['total_tokens_out']:,} out</div>
            </div>
            <div class="card">
                <div class="label">Haiku</div>
                <div class="value">{stats['calls_haiku']}</div>
                <div class="sub">{stats['tokens_in_haiku']:,} in / {stats['tokens_out_haiku']:,} out</div>
            </div>
            <div class="card">
                <div class="label">Sonnet</div>
                <div class="value">{stats['calls_sonnet']}</div>
                <div class="sub">{stats['tokens_in_sonnet']:,} in / {stats['tokens_out_sonnet']:,} out</div>
            </div>
            <div class="card">
                <div class="label">Admin (Opus)</div>
                <div class="value">{stats.get('calls_admin', 0)}</div>
                <div class="sub">{stats.get('tokens_in_admin', 0):,} in / {stats.get('tokens_out_admin', 0):,} out</div>
            </div>
        </div>
    </div>

    <div class="section">
        <h2>Estimated Cost</h2>
        <div class="grid">
            <div class="card">
                <div class="label">Total Cost</div>
                <div class="value cost">${costs['cost_total']:.4f}</div>
                <div class="sub">All models combined</div>
            </div>
            <div class="card">
                <div class="label">Haiku Cost</div>
                <div class="value cost">${costs['cost_haiku']:.4f}</div>
                <div class="sub">$1.00 / $5.00 per MTok</div>
            </div>
            <div class="card">
                <div class="label">Sonnet Cost</div>
                <div class="value cost">${costs['cost_sonnet']:.4f}</div>
                <div class="sub">$3.00 / $15.00 per MTok</div>
            </div>
            <div class="card">
                <div class="label">Admin Cost</div>
                <div class="value cost">${costs.get('cost_admin', 0.0):.4f}</div>
                <div class="sub">$15.00 / $75.00 per MTok (Opus)</div>
            </div>
        </div>
    </div>

    <div class="section">
        <h2>Recent Searches</h2>
        <table>
            <thead>
                <tr><th>Query</th><th>Results</th><th>Feedback</th><th>Date</th></tr>
            </thead>
            <tbody>
                {searches_html if searches_html else '<tr><td colspan="4">No searches yet</td></tr>'}
            </tbody>
        </table>
    </div>

    <div class="section" style="margin-top:40px; padding-top:16px; border-top:1px solid #e0e0e0; color:#888; font-size:12px;">
        FarmaFacil v{__version__} &bull;
        <a href="/api/v1/stats?phone={safe_phone}" style="color:#1a73e8;">JSON API</a>
    </div>
</body>
</html>"""
    return HTMLResponse(html)


# ── Scheduled Tasks API ───────────────────────────────────────────────


@router.get("/api/v1/scheduled-tasks")
@limiter.limit("60/minute")
async def list_scheduled_tasks(request: Request, _admin: str = Depends(_require_admin)) -> list[dict]:
    """List all scheduled tasks with their status."""
    from farmafacil.models.database import ScheduledTask

    async with async_session() as session:
        result = await session.execute(
            select(ScheduledTask).order_by(ScheduledTask.id)
        )
        tasks = result.scalars().all()

    return [
        {
            "id": t.id,
            "name": t.name,
            "task_key": t.task_key,
            "interval_minutes": t.interval_minutes,
            "enabled": t.enabled,
            "status": t.status,
            "last_run_at": t.last_run_at.isoformat() if t.last_run_at else None,
            "next_run_at": t.next_run_at.isoformat() if t.next_run_at else None,
            "last_result": t.last_result,
            "last_duration_seconds": t.last_duration_seconds,
        }
        for t in tasks
    ]


@router.post("/api/v1/scheduled-tasks/{task_id}/run")
@limiter.limit("10/minute")
async def run_scheduled_task(request: Request, task_id: int, _admin: str = Depends(_require_admin)) -> dict:
    """Manually trigger a scheduled task."""
    from farmafacil.services.scheduler import run_task_now

    result = await run_task_now(task_id)
    return {"task_id": task_id, "result": result}


# ── Conversation Session Viewer + CSV/DOCX Export ────────────────────

# A "session" is a continuous conversation — messages separated by
# gaps larger than this are considered separate sessions.
SESSION_GAP_MINUTES = 30


def _group_into_sessions(messages: list) -> list[dict]:
    """Group ordered messages into sessions by time gap.

    Returns:
        List of dicts: {start, end, messages: [...], count}
    """
    from datetime import timedelta

    if not messages:
        return []

    sessions = []
    current: list = [messages[0]]

    for i in range(1, len(messages)):
        prev_ts = messages[i - 1].created_at
        curr_ts = messages[i].created_at
        if prev_ts and curr_ts:
            gap = curr_ts - prev_ts
            if gap > timedelta(minutes=SESSION_GAP_MINUTES):
                sessions.append({
                    "start": current[0].created_at,
                    "end": current[-1].created_at,
                    "messages": current,
                    "count": len(current),
                })
                current = []
        current.append(messages[i])

    if current:
        sessions.append({
            "start": current[0].created_at,
            "end": current[-1].created_at,
            "messages": current,
            "count": len(current),
        })

    return sessions


@router.get("/admin/conversations", response_class=HTMLResponse)
@limiter.limit("60/minute")
async def admin_conversations_list(request: Request, _admin: str = Depends(_require_admin)) -> HTMLResponse:
    """Render a list of users with links to their conversation sessions."""
    async with async_session() as session:
        result = await session.execute(
            select(
                ConversationLog.phone_number,
                func.count(ConversationLog.id).label("msg_count"),
                func.max(ConversationLog.created_at).label("last_msg"),
            )
            .group_by(ConversationLog.phone_number)
            .order_by(func.max(ConversationLog.created_at).desc())
        )
        phones = result.all()

        user_result = await session.execute(select(User))
        users_by_phone = {
            u.phone_number: u.name or "(sin nombre)"
            for u in user_result.scalars().all()
        }

    rows = []
    for phone, count, last_msg in phones:
        name = users_by_phone.get(phone, "(desconocido)")
        safe_phone = escape(phone)
        safe_name = escape(name)
        last_str = last_msg.strftime("%Y-%m-%d %H:%M") if last_msg else ""
        rows.append(
            f'<tr>'
            f'<td><a href="/admin/conversations/{safe_phone}">{safe_name}</a></td>'
            f'<td>{safe_phone}</td><td>{count}</td><td>{last_str}</td>'
            f'<td>'
            f'<a href="/api/v1/conversations/export?phone={safe_phone}&format=csv">CSV</a> · '
            f'<a href="/api/v1/conversations/export?phone={safe_phone}&format=docx">Word</a>'
            f'</td></tr>'
        )

    html = f"""<!DOCTYPE html>
<html><head><title>Conversations</title>
<style>
body{{font-family:system-ui,sans-serif;max-width:1000px;margin:2em auto;padding:0 1em;}}
h1{{color:#1a73e8;}}
table{{width:100%;border-collapse:collapse;margin-top:1em;}}
th,td{{padding:.5em;text-align:left;border-bottom:1px solid #ddd;}}
th{{background:#f5f5f5;}}
a{{color:#1a73e8;text-decoration:none;}}
a:hover{{text-decoration:underline;}}
.nav{{margin-bottom:1em;}}
.nav a{{margin-right:1em;}}
</style></head><body>
<div class="nav">
  <a href="/admin">← Admin</a>
  <a href="/api/v1/conversations/export?format=csv">Export all (CSV)</a>
</div>
<h1>Conversations</h1>
<p>{len(rows)} users with conversation history. Click a name to see sessions.</p>
<table>
<tr><th>Name</th><th>Phone</th><th>Messages</th><th>Last</th><th>Export</th></tr>
{"".join(rows)}
</table>
</body></html>"""
    return HTMLResponse(html)


@router.get("/admin/conversations/{phone}", response_class=HTMLResponse)
@limiter.limit("60/minute")
async def admin_conversation_sessions(request: Request, phone: str, _admin: str = Depends(_require_admin)) -> HTMLResponse:
    """List all sessions for a user, each as a collapsible thread."""
    safe_phone = escape(phone)

    async with async_session() as session:
        user_result = await session.execute(
            select(User).where(User.phone_number == phone)
        )
        user = user_result.scalar_one_or_none()
        user_name = (user.name if user else None) or "(sin nombre)"

        msg_result = await session.execute(
            select(ConversationLog)
            .where(ConversationLog.phone_number == phone)
            .order_by(ConversationLog.created_at.asc())
        )
        messages = msg_result.scalars().all()

    sessions = _group_into_sessions(messages)

    session_html = []
    for idx, sess in enumerate(reversed(sessions), 1):  # newest first
        session_num = len(sessions) - idx + 1
        start = sess["start"].strftime("%Y-%m-%d %H:%M") if sess["start"] else "?"
        end = sess["end"].strftime("%H:%M") if sess["end"] else "?"
        duration_min = 0
        if sess["start"] and sess["end"]:
            duration_min = round((sess["end"] - sess["start"]).total_seconds() / 60)

        bubbles = []
        for msg in sess["messages"]:
            ts = msg.created_at.strftime("%H:%M:%S") if msg.created_at else ""
            is_inbound = msg.direction == "inbound"
            bubble_class = "inbound" if is_inbound else "outbound"
            label = "👤" if is_inbound else "🤖"
            if msg.message_type == "admin_out":
                label = "🛠️"
                bubble_class = "admin"
            text = escape(msg.message_text or "").replace("\n", "<br>")
            bubbles.append(
                f'<div class="msg {bubble_class}">'
                f'<div class="meta">{label} {ts}</div>'
                f'<div class="text">{text}</div>'
                f'</div>'
            )

        # Session ID is the start timestamp for export link
        session_ts = sess["start"].strftime("%Y%m%d_%H%M%S") if sess["start"] else f"session_{session_num}"
        session_iso = sess["start"].isoformat() if sess["start"] else ""

        session_html.append(f"""
<details {'open' if idx == 1 else ''} class="session">
<summary>
  <strong>Sesión {session_num}</strong> — {start} → {end}
  ({sess["count"]} mensajes, {duration_min} min)
  <span class="export-links">
    <a href="/api/v1/conversations/export?phone={safe_phone}&session={session_iso}&format=csv" onclick="event.stopPropagation()">CSV</a> ·
    <a href="/api/v1/conversations/export?phone={safe_phone}&session={session_iso}&format=docx" onclick="event.stopPropagation()">Word</a>
  </span>
</summary>
<div class="bubbles">
{"".join(bubbles)}
</div>
</details>
""")

    html = f"""<!DOCTYPE html>
<html><head><title>Sesiones — {escape(user_name)}</title>
<style>
body{{font-family:system-ui,sans-serif;max-width:900px;margin:2em auto;padding:0 1em;background:#f0f0f0;}}
h1{{color:#1a73e8;}}
.nav{{margin-bottom:1em;}}
.nav a{{margin-right:1em;color:#1a73e8;text-decoration:none;}}
.summary-box{{background:#fff;padding:1em;border-radius:8px;margin-bottom:1em;}}
.session{{background:#fff;border-radius:8px;margin:1em 0;padding:0;}}
.session summary{{padding:1em;cursor:pointer;font-size:1em;border-radius:8px;}}
.session summary:hover{{background:#f5f5f5;}}
.session[open] summary{{border-bottom:1px solid #ddd;border-radius:8px 8px 0 0;}}
.export-links{{float:right;font-size:.85em;}}
.export-links a{{color:#1a73e8;margin:0 .3em;}}
.bubbles{{padding:1em;}}
.msg{{margin:.6em 0;padding:.7em 1em;border-radius:12px;max-width:80%;word-wrap:break-word;}}
.msg.inbound{{background:#e8f0fe;border:1px solid #bbd0f7;margin-right:auto;}}
.msg.outbound{{background:#d4edda;border:1px solid #c3e6cb;margin-left:auto;}}
.msg.admin{{background:#fff3cd;border:1px solid #ffeaa7;margin-left:auto;}}
.meta{{font-size:.7em;color:#666;margin-bottom:.2em;}}
.text{{font-size:.95em;white-space:pre-wrap;}}
</style></head><body>
<div class="nav">
  <a href="/admin/conversations">← All users</a>
  <a href="/api/v1/conversations/export?phone={safe_phone}&format=csv">Export all sessions (CSV)</a>
  <a href="/api/v1/conversations/export?phone={safe_phone}&format=docx">Export all (Word)</a>
</div>
<div class="summary-box">
  <h1>{escape(user_name)}</h1>
  <p>Phone: <code>{safe_phone}</code> • {len(messages)} total messages • {len(sessions)} session(s)</p>
  <p style="color:#666;font-size:.85em;">Sessions group messages within {SESSION_GAP_MINUTES} minutes of each other.</p>
</div>
{"".join(session_html)}
</body></html>"""
    return HTMLResponse(html)


def _filter_session_messages(messages: list, session_iso: str) -> list:
    """Filter messages to the specific session starting at the given timestamp."""
    from datetime import datetime, timedelta

    try:
        target_start = datetime.fromisoformat(session_iso)
    except ValueError:
        return messages

    # Find the session that starts at or near target_start
    sessions = _group_into_sessions(messages)
    for sess in sessions:
        if sess["start"] and abs((sess["start"] - target_start).total_seconds()) < 2:
            return sess["messages"]

    return []


@router.get("/api/v1/conversations/export")
@limiter.limit("10/minute")
async def export_conversations(
    request: Request,
    phone: str | None = Query(None),
    session: str | None = Query(None, description="ISO timestamp of session start"),
    format: str = Query("csv", pattern="^(csv|docx)$"),
    _admin: str = Depends(_require_admin),
) -> StreamingResponse:
    """Export conversations as CSV or Word document.

    Args:
        phone: Optional — filter to a single user.
        session: Optional — ISO timestamp of a specific session's start.
        format: 'csv' or 'docx'.
    """
    async with async_session() as session_db:
        stmt = select(ConversationLog)
        if phone:
            stmt = stmt.where(ConversationLog.phone_number == phone)
        stmt = stmt.order_by(ConversationLog.created_at.asc())
        result = await session_db.execute(stmt)
        messages = result.scalars().all()

        user_result = await session_db.execute(select(User))
        users_by_phone = {
            u.phone_number: u.name or "" for u in user_result.scalars().all()
        }

    if session:
        messages = _filter_session_messages(messages, session)

    if format == "docx":
        return _export_as_docx(messages, users_by_phone, phone, session)

    return _export_as_csv(messages, users_by_phone, phone, session)


def _sanitize_filename_part(value: str | None) -> str:
    """Strip unsafe characters from a string used in Content-Disposition filenames.

    Prevents header injection via crafted phone/session parameters.
    (Item 61, v0.24.0.)
    """
    return re.sub(r"[^A-Za-z0-9_\-]", "_", value or "")[:30]


def _export_as_csv(messages, users_by_phone, phone, session_iso):
    """Produce a CSV export."""
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow([
        "id", "timestamp", "phone", "name", "direction",
        "message_type", "message_text", "wa_message_id",
    ])
    for msg in messages:
        writer.writerow([
            msg.id,
            msg.created_at.isoformat() if msg.created_at else "",
            msg.phone_number,
            users_by_phone.get(msg.phone_number, ""),
            msg.direction,
            msg.message_type,
            msg.message_text or "",
            msg.wa_message_id or "",
        ])

    buffer.seek(0)
    suffix = f"_{_sanitize_filename_part(phone)}" if phone else "_all"
    if session_iso:
        suffix += f"_session_{_sanitize_filename_part(session_iso)}"
    filename = f"conversations{suffix}.csv"

    return StreamingResponse(
        iter([buffer.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_as_docx(messages, users_by_phone, phone, session_iso):
    """Produce a Word document export with chat-transcript styling."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = "FarmaFacil Conversation"
    if phone:
        name = users_by_phone.get(phone, "")
        title += f" — {name} ({phone})" if name else f" — {phone}"
    doc.add_heading(title, level=1)

    if session_iso:
        doc.add_paragraph(f"Session: {session_iso}").italic = True

    doc.add_paragraph(f"Total messages: {len(messages)}").italic = True
    doc.add_paragraph("")

    for msg in messages:
        ts = msg.created_at.strftime("%Y-%m-%d %H:%M:%S") if msg.created_at else ""
        is_inbound = msg.direction == "inbound"
        name = users_by_phone.get(msg.phone_number, "")

        if msg.message_type == "admin_out":
            label = "🛠️ Admin AI"
            color = RGBColor(0xB3, 0x7F, 0x00)  # amber
        elif is_inbound:
            label = f"👤 {name or msg.phone_number}"
            color = RGBColor(0x1A, 0x73, 0xE8)  # blue
        else:
            label = "🤖 FarmaFacil Bot"
            color = RGBColor(0x2E, 0x7D, 0x32)  # green

        # Header paragraph
        header = doc.add_paragraph()
        header_run = header.add_run(f"{label} — {ts}")
        header_run.bold = True
        header_run.font.color.rgb = color
        header_run.font.size = Pt(9)

        # Message body
        body = doc.add_paragraph(msg.message_text or "")
        for run in body.runs:
            run.font.size = Pt(11)

        doc.add_paragraph("")  # spacer

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    suffix = f"_{_sanitize_filename_part(phone)}" if phone else "_all"
    if session_iso:
        suffix += f"_session_{_sanitize_filename_part(session_iso)}"
    filename = f"conversations{suffix}.docx"

    return StreamingResponse(
        iter([buffer.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Voice message audio playback ───────────────────────────────────────


@router.get("/api/v1/audio/{voice_message_id}")
@limiter.limit("30/minute")
async def get_voice_audio(
    voice_message_id: int,
    request: Request,
    _admin: str = Depends(_require_admin),
) -> FileResponse:
    """Serve a stored voice message audio file for playback.

    Protected by HTTP Basic auth (admin credentials). Used by the
    SQLAdmin dashboard to embed an HTML5 ``<audio>`` player.

    Args:
        voice_message_id: The voice_messages.id to serve.
        request: The incoming HTTP request.

    Returns:
        The audio file with appropriate MIME type.
    """
    from farmafacil.services.voice import AUDIO_BASE_DIR, get_audio_absolute_path

    async with async_session() as session:
        stmt = select(VoiceMessage).where(VoiceMessage.id == voice_message_id)
        voice_msg = (await session.execute(stmt)).scalar_one_or_none()

    if voice_msg is None:
        raise HTTPException(status_code=404, detail="Voice message not found")

    audio_path = get_audio_absolute_path(voice_msg.audio_path).resolve()

    # Path containment: audio files must reside under AUDIO_BASE_DIR
    try:
        audio_path.relative_to(AUDIO_BASE_DIR.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Invalid audio path")

    if not audio_path.exists():
        raise HTTPException(status_code=404, detail="Audio file not found on disk")

    # Determine MIME type from extension
    suffix = audio_path.suffix.lower()
    mime_map = {
        ".ogg": "audio/ogg",
        ".opus": "audio/opus",
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".wav": "audio/wav",
    }
    media_type = mime_map.get(suffix, "audio/ogg")

    return FileResponse(
        path=str(audio_path),
        media_type=media_type,
        filename=audio_path.name,
    )
